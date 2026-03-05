"""
docx_builder.py — Converts plain-text rewritten resume into a formatted DOCX.
Uses python-docx. Returns bytes ready for streaming.
"""

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color palette ─────────────────────────────────────────────────────────────
INK    = RGBColor(0x14, 0x14, 0x1F)   # near-black
ACCENT = RGBColor(0xC9, 0xA2, 0x27)   # gold
GRAY   = RGBColor(0x44, 0x44, 0x55)   # dark gray
LGRAY  = RGBColor(0x88, 0x88, 0x99)   # light gray


def _set_font(run, name="Calibri", size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_para(doc, text="", style="Normal"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _hr(doc, color="C9A227"):
    """Add a thin colored horizontal rule via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def build_resume_docx(plain_text: str, candidate_name: str = "Candidate", role_title: str = "") -> bytes:
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

    # ── Name header ──────────────────────────────────────────────────
    name_p = _add_para(doc)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(candidate_name.upper())
    _set_font(r, "Calibri", 22, bold=True, color=INK)

    if role_title:
        role_p = _add_para(doc)
        role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_p.paragraph_format.space_after = Pt(6)
        r2 = role_p.add_run(role_title)
        _set_font(r2, "Calibri", 11, bold=False, color=ACCENT)

    _hr(doc)

    # ── Parse sections ────────────────────────────────────────────────
    # Section headers are lines followed by "---..." underlines
    lines = plain_text.splitlines()
    i = 0
    current_section = None

    while i < len(lines):
        line = lines[i].rstrip()

        # Detect section header: next line is "---..." dashes
        is_header = (
            i + 1 < len(lines) and
            re.match(r"^-{3,}$", lines[i + 1].strip()) and
            line.strip()
        )

        if is_header:
            current_section = line.strip().upper()
            # Section heading
            h = _add_para(doc)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after  = Pt(1)
            hr = h.add_run(current_section)
            _set_font(hr, "Calibri", 10, bold=True, color=ACCENT)
            _hr(doc, "C9A227")
            i += 2   # skip header + dashes line
            continue

        # Bullet point
        if line.startswith("- "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.2)
            r = p.add_run(bullet_text)
            _set_font(r, "Calibri", 10, color=GRAY)
            i += 1
            continue

        # Experience / project sub-header (contains | separators or — dash)
        if ("|" in line or " — " in line or " - " in line) and line.strip() and current_section in ("EXPERIENCE", "PROJECTS"):
            p = _add_para(doc)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(line.strip())
            _set_font(r, "Calibri", 10, bold=True, color=INK)
            i += 1
            continue

        # Skills line (comma-separated, no bullets)
        if current_section == "TECHNICAL SKILLS" and line.strip():
            p = _add_para(doc)
            p.paragraph_format.space_after = Pt(2)
            # Bold the skill names (before comma), normal for rest
            skills = [s.strip() for s in line.split(",") if s.strip()]
            for j, skill in enumerate(skills):
                r = p.add_run(skill)
                _set_font(r, "Calibri", 10, color=GRAY)
                if j < len(skills) - 1:
                    sep = p.add_run("  ·  ")
                    _set_font(sep, "Calibri", 10, color=LGRAY)
            i += 1
            continue

        # Normal paragraph
        if line.strip():
            p = _add_para(doc)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line.strip())
            _set_font(r, "Calibri", 10, color=GRAY)

        i += 1

    # ── Serialize to bytes ────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()