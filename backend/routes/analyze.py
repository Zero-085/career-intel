from fastapi import APIRouter, UploadFile, File, Form
from services.llm_adapter import generate_analysis
from utils.prompt_builder import build_prompt
from utils.json_parser import extract_json
from utils.score_enforcer import enforce_scores
from utils.jd_scraper import scrape_jd

import pdfplumber
import docx
import io
import logging

logger = logging.getLogger(__name__)
router = APIRouter()


def _extract_resume_text(resume_file, resume_text):
    """Extract resume text from file upload or return raw text."""
    if resume_file:
        content = resume_file  # already bytes from caller
        if resume_file.filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif resume_file.filename.endswith(".docx"):
            d = docx.Document(io.BytesIO(content))
            return "\n".join(p.text for p in d.paragraphs)
        else:
            return None, "Unsupported file format. Upload PDF or DOCX."
    return resume_text, None


@router.post("/analyze")
async def analyze(
    resume_text: str = Form(None),
    jd_text:     str = Form(None),
    jd_url:      str = Form(None),
    resume_file: UploadFile = File(None),
):
    # ── 1. Resume text ────────────────────────────────────────────────
    if resume_file:
        content = await resume_file.read()
        if resume_file.filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                resume_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif resume_file.filename.endswith(".docx"):
            d = docx.Document(io.BytesIO(content))
            resume_text = "\n".join(p.text for p in d.paragraphs)
        else:
            return {"error": "Unsupported file format. Upload PDF or DOCX."}

    if not resume_text or not resume_text.strip():
        return {"error": "Resume text is required."}

    # ── Detect if user pasted JD in the resume field ──────────────────
    jd_phrases = [
        "we are hiring", "we are looking for", "about the role",
        "required skills:", "preferred skills:", "responsibilities:",
        "what you will do", "what we're looking for", "job description",
        "you will be responsible", "the ideal candidate", "join our team",
    ]
    personal_signals = [
        "my experience", "i have", "i am ", "i've", "my name",
        "summary:", "objective:", "worked at", "i built", "i led",
    ]
    resume_lower = resume_text.lower()[:800]
    jd_hits      = sum(1 for p in jd_phrases if p in resume_lower)
    has_personal = any(p in resume_lower for p in personal_signals)

    if jd_hits >= 2 and not has_personal:
        return {
            "error": "It looks like you pasted the job description in the Resume field. Please put your personal resume (name, experience, skills) in the left box, and the job description in the right box.",
            "hint": "resume_jd_swap",
        }

    # ── 2. JD text (URL or paste) ─────────────────────────────────────
    scraped_meta = {}
    if jd_url and jd_url.strip():
        result = scrape_jd(jd_url.strip())
        if not result["success"]:
            return {
                "error": result["error"],
                "blocked": result.get("blocked", False),
                "hint": "Paste the job description text manually instead.",
            }
        jd_text = result["text"]
        scraped_meta = {
            "jd_title":    result.get("title", ""),
            "jd_company":  result.get("company", ""),
            "jd_platform": result.get("platform", ""),
        }
    elif not jd_text or not jd_text.strip():
        return {"error": "Job description is required (text or URL)."}

    # ── 3. Build prompt & call LLM ────────────────────────────────────
    prompt = build_prompt(resume_text, jd_text)
    try:
        raw = generate_analysis(prompt)
    except Exception as e:
        logger.error(f"LLM call failed: {e}")
        return {"error": f"LLM service error: {str(e)}"}

    if not raw:
        return {"error": "LLM returned an empty response."}

    # ── 4. Parse & enforce ────────────────────────────────────────────
    parsed = extract_json(raw)
    if not parsed:
        logger.error(f"JSON parse failed. Raw (first 800):\n{raw[:800]}")
        return {
            "error": "Failed to parse LLM response.",
            "raw_preview": raw[:600],
            "hint": "The LLM may have returned truncated JSON. Try again.",
        }

    # ── Cross-check: reconcile scoring_breakdown vs skill_analysis ──────────
    # LLM sometimes under-counts required_matched in scoring_breakdown.
    # The matched_skills list is more reliable — use whichever count is higher.
    _bd   = parsed.get("scoring_breakdown") or {}
    _sa   = parsed.get("skill_analysis") or {}
    _mat  = [s for s in (_sa.get("matched_skills") or []) if s]
    _mis  = _sa.get("missing_skills") or []
    _n_matched = len(_mat)
    _n_total   = _n_matched + len(_mis)
    _bd_matched = int(_bd.get("required_matched", 0))
    _bd_total   = int(_bd.get("required_total", 0))
    if _n_matched > _bd_matched and _n_total > 0:
        _bd["required_matched"] = _n_matched
        if _bd_total == 0 or _n_total > _bd_total:
            _bd["required_total"] = _n_total
        parsed["scoring_breakdown"] = _bd

    result = enforce_scores(parsed)

    # Attach scraped metadata if URL was used
    if scraped_meta:
        result.update(scraped_meta)

    # Always return resume_text so frontend can pass it to /rewrite-resume
    result["_resume_text"] = resume_text
    result["_jd_text"]     = jd_text

    return result


@router.post("/analyze/debug")
async def analyze_debug(
    resume_text: str = Form(None),
    jd_text:     str = Form(None),
    jd_url:      str = Form(None),
    resume_file: UploadFile = File(None),
):
    if resume_file:
        content = await resume_file.read()
        if resume_file.filename.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                resume_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif resume_file.filename.endswith(".docx"):
            d = docx.Document(io.BytesIO(content))
            resume_text = "\n".join(p.text for p in d.paragraphs)

    if jd_url and jd_url.strip() and not jd_text:
        scraped = scrape_jd(jd_url.strip())
        if scraped["success"]:
            jd_text = scraped["text"]

    if not resume_text:
        return {"error": "Resume text required."}

    prompt = build_prompt(resume_text, jd_text or "")
    raw    = generate_analysis(prompt)
    parsed = extract_json(raw)
    enforced = enforce_scores(parsed) if parsed else None

    return {
        "raw_response": raw,
        "parse_success": parsed is not None,
        "enforced_scores": {
            "match_score":       enforced.get("match_score") if enforced else None,
            "ats_optimization":  enforced.get("ats_optimization") if enforced else None,
            "scoring_breakdown": enforced.get("scoring_breakdown") if enforced else None,
        } if enforced else None,
    }


@router.post("/jd/fetch-url")
async def fetch_jd_url(jd_url: str = Form(...)):
    """Standalone endpoint to preview a scraped JD before analyzing."""
    result = scrape_jd(jd_url.strip())
    return result